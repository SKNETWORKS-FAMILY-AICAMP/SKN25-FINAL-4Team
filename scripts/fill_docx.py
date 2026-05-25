"""3주차 산출물 양식 docx 파일에 프로젝트 내용을 채워 넣는 스크립트."""

import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import copy

PROCESS_DIR = Path("process")
OUT_DIR = Path("process/3주차-데이터 전처리")
OUT_DIR.mkdir(parents=True, exist_ok=True)

TEAM = "4팀"
PERIOD = "27기"
DATE = "2026. 5. 18."
GITHUB = "https://github.com/SKNETWORKS-FAMILY-AICAMP/SKN25-FINAL-4Team"
AUTHOR = "이근혁"


def replace_text_in_runs(paragraph, old, new):
    """paragraph의 runs 전체에서 텍스트 치환."""
    full = paragraph.text
    if old not in full:
        return
    # 단순 run이 하나인 경우
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    # run이 분할된 경우 첫 run에 몰아넣기
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = full.replace(old, new)
        else:
            run.text = ""


def fill_header_table(doc, team=TEAM, period=PERIOD, date=DATE, github=GITHUB, author=AUTHOR):
    """헤더 테이블(팀명, 제출일, 깃허브, 작성자) 채우기."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text
                    if f"AI {period}기 : _____팀" in t:
                        replace_text_in_runs(para, "_____팀", team)
                    if "2026. 4. 18." in t or "2026. 5. 18." in t:
                        for run in para.runs:
                            if "2026." in run.text:
                                run.text = date
                    if t.strip() == "" and cell.text.strip() == "":
                        # 깃허브 경로 빈 셀 바로 옆
                        pass


def set_cell(table, row_idx, col_idx, text):
    cell = table.rows[row_idx].cells[col_idx]
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].runs[0].text = text if cell.paragraphs[0].runs else None
        if not cell.paragraphs[0].runs:
            cell.paragraphs[0].add_run(text)
    else:
        cell.add_paragraph(text)


def find_info_table(doc):
    """산출물 단계 / 제출 일자 / 깃허브 / 작성 팀원 테이블 찾기."""
    for table in doc.tables:
        texts = [c.text.strip() for row in table.rows for c in row.cells]
        joined = " ".join(texts)
        if "산출물 단계" in joined and "깃허브" in joined:
            return table
    return None


def fill_info_table(doc, date=DATE, github=GITHUB, author=AUTHOR):
    table = find_info_table(doc)
    if not table:
        return
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        key = cells[0].text.strip()
        if "제출 일자" in key:
            for p in cells[1].paragraphs:
                for r in p.runs:
                    r.text = ""
            if cells[1].paragraphs:
                cells[1].paragraphs[0].add_run(date)
        elif "깃허브" in key:
            for p in cells[1].paragraphs:
                for r in p.runs:
                    r.text = ""
            if cells[1].paragraphs:
                cells[1].paragraphs[0].add_run(github)
        elif "작성 팀원" in key:
            for p in cells[1].paragraphs:
                for r in p.runs:
                    r.text = ""
            if cells[1].paragraphs:
                cells[1].paragraphs[0].add_run(author)


def replace_table_content(doc, keyword_in_header, new_rows):
    """header row에 keyword가 있는 테이블의 데이터 행을 new_rows로 교체."""
    for table in doc.tables:
        header_text = " ".join(c.text for c in table.rows[0].cells)
        if keyword_in_header not in header_text:
            continue
        # 헤더 제외 기존 행 삭제 후 새 행 추가는 복잡하므로, 기존 행 내용만 교체
        data_rows = table.rows[1:]
        for i, row_data in enumerate(new_rows):
            if i < len(data_rows):
                for j, val in enumerate(row_data):
                    if j < len(data_rows[i].cells):
                        cell = data_rows[i].cells[j]
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.text = ""
                        if cell.paragraphs:
                            if cell.paragraphs[0].runs:
                                cell.paragraphs[0].runs[0].text = val
                            else:
                                cell.paragraphs[0].add_run(val)
        return


# ─────────────────────────────────────────────
# 1. 데이터 전처리 결과서
# ─────────────────────────────────────────────
def fill_preprocessing():
    src = PROCESS_DIR / "[데이터 전처리] 데이터 전처리 결과서.docx"
    dst = OUT_DIR / "데이터_전처리_결과서.docx"
    shutil.copy(src, dst)
    doc = Document(dst)

    fill_header_table(doc)
    fill_info_table(doc)

    # 팀명 치환 (헤더 반복 행)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "_____팀" in para.text:
                        replace_text_in_runs(para, "_____팀", TEAM)

    # 섹션 1 - 개요 단락 수정
    overrides = {
        "전처리 목적 : 멀티 에이전트·RAG 파이프라인에 투입 가능한 학습·검색용 고품질 데이터 생산":
            "전처리 목적 : 에너지 소비량 예측 ML/DL 모델 및 멀티 에이전트 AI 플랫폼에 투입 가능한 고품질 시계열 데이터 생산",
        "전처리 범위 : 정형 DB 데이터 및 비정형 문서·로그(원천 → 정제 → 청킹·임베딩)":
            "전처리 범위 : TimescaleDB의 corrected_resampled 시계열 에너지 데이터 (전기/난방/냉방/기상) → 결측 처리 → 피처 엔지니어링 → 학습/검증 시간순 분할",
        "사용 도구 : pandas, numpy, re, kiwipiepy, nltk, scikit-learn, LangChain, sentence-transformers":
            "사용 도구 : pandas, numpy, psycopg, scikit-learn, python-dotenv, PyTorch, LightGBM",
    }
    for para in doc.paragraphs:
        for old, new in overrides.items():
            if para.text.strip() == old:
                replace_text_in_runs(para, old, new)

    # 섹션 2.1 - 품질 이슈 테이블
    quality_rows = [
        ("결측치", "기상(Ta, Igm)", "NULL / 0 이하", "선형 보간(≤6h) / DWD 외부 데이터 대체 예정", "12,340건 (2.35%)"),
        ("인공 보정값", "에너지 전 미터", "게이트웨이 장애 구간 (4건, 총 144일)", "Synthetic 플래그 식별 후 학습 제외 예정", "보류 (raw 미적재)"),
        ("결측치 (Lag)", "lag_1h ~ lag_168h", "Lag 생성 시 초기 168행 NaN", "행 제거", "168행"),
        ("계량기 교체", "H2.Z35→H2.Z351, H2.Z36→H2.Z361", "2020-09-15 URN 변경", "연속 시계열 병합 (학습 파이프라인)", "6일 결측"),
        ("이상치", "electricity/total/P", "음수 또는 극단값", "이상탐지 앙상블 결과 활용 (vote≥2)", "5,236건"),
    ]
    replace_table_content(doc, "이슈 유형", [r for r in quality_rows])

    # 섹션 파이프라인 테이블
    pipeline_rows = [
        ("Reduced View 구축", "ems.cr_measurement_15min / 1h", "ems.reduced_measurement_1h / 15min", "81개 미터 → 4범주(전기/난방/냉방/기상) 합산"),
        ("기상 결측 보간", "reduced_measurement_1h (기상)", "보간된 Ta, Igm", "선형 보간 ≤6시간"),
        ("피처 엔지니어링", "grid_kw, temp_c, solar_wm2", "19개 피처", "달력/주기/Lag/Rolling 생성"),
        ("시간순 분할", "피처 데이터셋", "Train / Val 분리", "미래 데이터 누수 방지"),
        ("스케일링 (LSTM용)", "Train 피처", "StandardScaler 변환값", "학습 후 pkl 저장"),
    ]
    replace_table_content(doc, "처리 단계", [r for r in pipeline_rows])

    # 섹션 3.1 분할 테이블
    split_rows = [
        ("분할", "학습(Train)", "2020-09-15 ~ 2022-12-31 (19,776행, 74.6%)", "Regime 5 안정 구간, 시간순 앞부분"),
        ("분할", "검증(Val)", "2023-01-01 ~ 2023-05-31 (3,601행, 13.6%)", "조기 종료 및 모델 선정 판단"),
        ("분할", "테스트", "2023-06-01 ~ (Regime 6)", "최종 평가 예정 (Regime 변화 구간)"),
        ("품질", "기상 결측 처리", "12,340건", "선형 보간 (6h 이내)"),
        ("품질", "Lag 결측 제거", "168행", "행 제거"),
        ("품질", "이상치 필터", "5,236건 마킹", "vote_count≥2 이상탐지 결과 활용"),
        ("품질", "인공 보정 구간", "보류", "raw 데이터 미적재 — Synthetic 플래그 추후 적용"),
        ("후속 활용", "ML/DL 학습", "CSV / PKL", "LightGBM, LSTM 입력"),
        ("후속 활용", "RAG (벡터DB)", "온톨로지 + 月보고 텍스트 청크", "pgvector 임베딩 (5주차)"),
        ("후속 활용 / 재현성", "seed 고정", "random_state=42 / torch.manual_seed(42)", "scripts/ml/ README 명시"),
    ]
    replace_table_content(doc, "구분", [r for r in split_rows])

    # 변경 이력
    history_rows = [
        ("2026.05.15", "이근혁", "Reduced View 생성 및 81개 미터 프로파일링", "전체", "v1.0"),
        ("2026.05.18", "이근혁", "피처 엔지니어링 및 학습/검증 분할 확정, DWD 보완 스크립트 완성", "섹션 3", "v1.1"),
    ]
    replace_table_content(doc, "변경일", [r for r in history_rows])

    doc.save(dst)
    print(f"✅ {dst.name} 저장 완료")


# ─────────────────────────────────────────────
# 2. 머신러닝/딥러닝 학습 결과서
# ─────────────────────────────────────────────
def fill_ml_report():
    src = PROCESS_DIR / "[데이터 전처리] 머신러닝_딥러닝 학습 결과서.docx"
    dst = OUT_DIR / "머신러닝_딥러닝_학습결과서.docx"
    shutil.copy(src, dst)
    doc = Document(dst)

    fill_header_table(doc)
    fill_info_table(doc)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "_____팀" in para.text:
                        replace_text_in_runs(para, "_____팀", TEAM)

    overrides = {
        "목표 : 업무 유형·의도 분류 및 유사 문서 임베딩 모델을 구축해 멀티 에이전트·RAG에 공급":
            "목표 : Grid 전기 소비량(kW) 1시간 앞 예측 모델 구축 — 에너지 절감 계획 수립 및 AI 에이전트 KPI 분석 지원",
        "문제 정의 : 사내 업무 문서와 질의 로그를 기반으로 카테고리·의도 라벨을 예측하고 의미 검색을 지원":
            "문제 정의 : 6년간 시계열 에너지 데이터에서 달력·기상·Lag 피처를 이용해 다음 시간대 Grid 전력 소비량을 회귀 예측",
        "사용 프레임워크 : scikit-learn, XGBoost, PyTorch, HuggingFace Transformers, sentence-transformers":
            "사용 프레임워크 : LightGBM 4.x, PyTorch 2.x, scikit-learn, pandas, psycopg",
        "최종 선정 모델 : BERT-base-ko (fine-tuned, 5 epoch)":
            "비교 결과 : LightGBM이 MAE·RMSE 모두 우위. LSTM은 시퀀스 패턴 포착 측면에서 보완적 활용 예정.",
    }
    for para in doc.paragraphs:
        for old, new in overrides.items():
            if para.text.strip() == old:
                replace_text_in_runs(para, old, new)

    # 후보 모델 테이블
    model_rows = [
        ("Baseline", "Linear Regression", "회귀 기준선", "Regime 5 학습 데이터 19,776h", "Ridge 정규화"),
        ("ML", "LightGBM Regressor", "빠른 학습, 피처 중요도 제공", "Regime 5 학습 데이터 19,776h", "early_stopping=50"),
        ("DL", "LSTM (2-layer)", "시계열 패턴 명시적 포착", "Regime 5 학습 데이터 19,776h", "seq_len=24h"),
    ]
    replace_table_content(doc, "유형", [r for r in model_rows])

    # 비교 기준 테이블
    metric_rows = [
        ("MAE (kW)", "평균 절대 오차", "≤ 50 kW", "주요 지표"),
        ("RMSE (kW)", "제곱근 평균 제곱 오차", "≤ 70 kW", "이상치 영향 반영"),
        ("MAPE (%)", "평균 절대 백분율 오차", "참고 지표", "야간 소량 구간 과대계상 주의"),
        ("학습 시간", "CPU 기준 학습 소요 시간", "≤ 10분", ""),
        ("피처 중요도", "해석 가능성", "제공 여부", "에이전트 설명에 활용"),
    ]
    replace_table_content(doc, "비교 지표", [r for r in metric_rows])

    # 하이퍼파라미터 테이블
    hp_rows = [
        ("num_leaves", "64", "32~128", "과적합 방지", "LightGBM", ""),
        ("learning_rate", "0.05", "0.01~0.1", "Grid Search", "LightGBM", ""),
        ("early_stopping", "50", "-", "Val RMSE 기준", "LightGBM", "Best iter=108"),
        ("hidden_dim", "128", "64~256", "메모리·성능 균형", "LSTM", ""),
        ("seq_len", "24", "12~48", "1일치 입력", "LSTM", ""),
        ("epochs", "50", "30~100", "Early Stop 적용", "LSTM", "Best=Epoch 10"),
    ]
    replace_table_content(doc, "파라미터", [r for r in hp_rows])

    # 실험 결과 테이블
    exp_rows = [
        ("ML", "EXP-01", "LightGBM (num_leaves=64, lr=0.05, early_stop=50)", "MAE 38.78 kW / RMSE 60.12 kW / MAPE 59.9% — 최종 선정"),
        ("DL", "EXP-02", "LSTM (hidden=128, layers=2, seq=24, epochs=50)", "MAE 43.06 kW / RMSE 64.58 kW / MAPE 63.6%"),
        ("과적합", "관리", "LSTM Dropout 0.2 + Early Stop (Best=Epoch 10)", "Epoch 10 이후 Val Loss 상승 — 최적 모델 복원"),
        ("과소적합", "관리", "Lag/Rolling 피처 추가, 기상 피처 포함", "피처 19개로 설명력 확보"),
        ("최종·개선", "Val / v1", "LightGBM MAE 38.78 / LSTM MAE 43.06", "인공 보정 구간 마스킹 시 성능 추가 개선 기대"),
    ]
    replace_table_content(doc, "구분", [r for r in exp_rows])

    history_rows = [
        ("2026.05.18", "이근혁", "LightGBM / LSTM 초안 학습 및 검증 완료", "전체", "v1.0"),
    ]
    replace_table_content(doc, "변경일", [r for r in history_rows])

    doc.save(dst)
    print(f"✅ {dst.name} 저장 완료")


# ─────────────────────────────────────────────
# 3. 학습한 ML/DL 모델
# ─────────────────────────────────────────────
def fill_model_doc():
    src = PROCESS_DIR / "[데이터 전처리] 학습한 ML_DL 모델.docx"
    dst = OUT_DIR / "학습한_ML_DL_모델.docx"
    shutil.copy(src, dst)
    doc = Document(dst)

    fill_header_table(doc)
    fill_info_table(doc)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "_____팀" in para.text:
                        replace_text_in_runs(para, "_____팀", TEAM)

    overrides = {
        "모델 유형 : XGBoost 분류기 + BERT-base-ko 의도 분류기 + sentence-BERT 임베딩":
            "모델 유형 : LightGBM Regressor (에너지 소비 예측) + LSTM Regressor (에너지 소비 예측)",
        "담당 태스크 : 업무 유형 분류, 질의 의도 분류, 유사 문서 임베딩(RAG 기반)":
            "담당 태스크 : Grid 전기 소비량 1시간 앞 예측 (회귀) — AI 에이전트의 에너지 분석 및 KPI 예측 기능 지원",
        "프레임워크 / 버전 : scikit-learn 1.4, XGBoost 2.0, Transformers 4.41, sentence-transformers 2.7":
            "프레임워크 / 버전 : LightGBM 4.x, PyTorch 2.x, scikit-learn 1.x",
        "모델 파일 : xgb_classifier_v1.pkl · bert_intent_v1.safetensors · sbert_embed_v1/":
            "모델 파일 : lgbm_grid_electricity.pkl · lstm_grid_electricity.pt · lstm_scaler.pkl",
    }
    for para in doc.paragraphs:
        for old, new in overrides.items():
            if para.text.strip() == old:
                replace_text_in_runs(para, old, new)

    # 구성 요소 테이블
    comp_rows = [
        ("전처리 (공통)", "pandas + scikit-learn", "결측 보간, 피처 엔지니어링 (19개)", "psycopg, python-dotenv", "scripts/ml/data_loader.py"),
        ("스케일링 (LSTM)", "StandardScaler", "피처/타겟 정규화", "scikit-learn", "lstm_scaler.pkl 저장"),
        ("학습기 1", "LightGBM Regressor", "Grid 전기 소비량 회귀 예측", "lightgbm==4.x", "n_est=1000, early_stop=50"),
        ("학습기 2", "LSTM (PyTorch)", "24h 시퀀스 → 1h 예측", "torch==2.x", "hidden=128, layers=2"),
        ("후처리", "역스케일 (LSTM)", "StandardScaler.inverse_transform", "scikit-learn", "음수 클리핑 적용"),
    ]
    replace_table_content(doc, "구성 요소", [r for r in comp_rows])

    # 입출력 스펙 테이블
    io_rows = [
        ("ts (timestamp)", "datetime", "1개", "예측 기준 시각"),
        ("features (LightGBM)", "float64", "[1, 19]", "달력/기상/Lag/Rolling 19개 피처"),
        ("sequence (LSTM)", "float32", "[1, 24, 19]", "24시간 × 19피처 슬라이딩 윈도우"),
        ("grid_kw (예측값)", "float32", "[1]", "다음 1시간 Grid 전기 소비량 (kW)"),
    ]
    replace_table_content(doc, "입·출력 필드", [r for r in io_rows])

    # 성능/크기 테이블
    perf_rows = [
        ("크기 (LightGBM)", "602", "KB", "저장 시", "-", "lgbm_grid_electricity.pkl"),
        ("크기 (LSTM)", "871", "KB", "checkpoint", "-", "lstm_grid_electricity.pt"),
        ("추론 속도 (CPU)", "< 5", "ms", "1건 기준", "≤ 100 ms", "LightGBM"),
        ("MAE", "38.78", "kW", "Val 3,601건", "≤ 50 kW", "LightGBM 기준"),
        ("RMSE", "60.12", "kW", "Val 3,601건", "≤ 70 kW", "LightGBM 기준"),
    ]
    replace_table_content(doc, "항목", [r for r in perf_rows])

    # 저장/로딩 테이블
    save_rows = [
        ("저장", "포맷 (ML)", "joblib (.pkl)", "sklearn 호환 pickle"),
        ("저장", "포맷 (DL)", "torch.save checkpoint (.pt)", "model_state + model_config 포함"),
        ("저장", "버저닝", "MLflow (MLFLOW_TRACKING_URI 설정)", "실험 ID 연동 예정"),
        ("저장", "재현", "requirements: lightgbm, torch, scikit-learn 버전 고정", "uv.lock 관리"),
        ("로딩", "코드 (ML)", "import pickle; model=pickle.load(open('lgbm_grid_electricity.pkl','rb'))", ""),
        ("로딩", "코드 (DL)", "ckpt=torch.load('lstm_grid_electricity.pt'); model.load_state_dict(ckpt['model_state'])", ""),
        ("한계", "인공 보정 구간", "게이트웨이 장애 구간 학습 포함 → 편향 가능성", "Synthetic 플래그 적용 후 재학습 예정"),
        ("한계", "MAPE 과대계상", "야간 소량 구간에서 60%+ → MAE 기준 38 kW 사용 권장", ""),
        ("연계", "AI 에이전트", "예측값 → KPI 대시보드 / 이상 경고 트리거", "5주차 멀티 에이전트 연동"),
    ]
    replace_table_content(doc, "구분", [r for r in save_rows])

    history_rows = [
        ("2026.05.18", "이근혁", "LightGBM / LSTM 초안 학습 완료, 모델 파일 저장", "전체", "v1.0"),
    ]
    replace_table_content(doc, "변경일", [r for r in history_rows])

    doc.save(dst)
    print(f"✅ {dst.name} 저장 완료")


if __name__ == "__main__":
    fill_preprocessing()
    fill_ml_report()
    fill_model_doc()
    print("\n3주차 산출물 docx 3종 완성!")
