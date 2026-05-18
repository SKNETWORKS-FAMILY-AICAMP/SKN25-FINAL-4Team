"""81개 개별 계량기 상세 분석 Word(DOCX) 보고서 생성 스크립트.

가독성을 높이기 위해 python-docx를 사용하여 여백을 최소화하고,
통계 기반의 자동 분석 요약(Insight)을 표 아래에 삽입하며,
각 계량기마다 정확히 1페이지씩 출력되도록 제어합니다.

Usage:
    uv run --with pandas --with python-docx python scripts/profiling/generate_individual_meter_report.py
"""

import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 경로 설정
ROOT_DIR = Path(__file__).parent.parent.parent
PROFILING_DIR = ROOT_DIR / "outputs" / "profiling"
FIGURES_DIR = ROOT_DIR / "outputs" / "figures" / "meters" / "individual"
DOCS_DIR = ROOT_DIR / "docs"


def analyze_meter(meter_stats: pd.DataFrame, group: str) -> list[str]:
    """통계량을 바탕으로 요약 인사이트 텍스트를 생성합니다."""
    insights = []
    
    # 해당 장비의 역할
    insights.append(f"▶ 이 계량기는 [{group}] 그룹에 속하는 설비입니다.")
    
    # 전력(P) 기준으로 특성 파악
    p_stat = meter_stats[meter_stats["measurement"] == "P"]
    if not p_stat.empty:
        stat = p_stat.iloc[0]
        zero_ratio = stat.get("zero_ratio", 0)
        mean_v = stat.get("mean", 0)
        max_v = stat.get("max", 0)
        
        # 1. 운영 패턴 분석
        if zero_ratio > 0.05:
            insights.append(f"▶ 무부하(0) 비율이 {zero_ratio*100:.1f}%로, 근무 외 시간(야간/휴일 등)에 가동을 멈추는 간헐적 운영 패턴을 보입니다.")
        else:
            insights.append(f"▶ 무부하 비율이 거의 0%에 가까워, 24시간 내내 일정하게 전력을 소모하는 상시 운영(기저부하) 패턴을 띱니다.")
            
        # 2. 부하 변동성 분석
        if mean_v > 0 and max_v > mean_v * 3:
            insights.append(f"▶ 평균 사용량({mean_v:.1f}) 대비 최대 피크({max_v:.1f})가 매우 높은 변동성 부하 특성이 관측됩니다. 피크 관리에 유의해야 합니다.")
            
    return insights


def generate_docx():
    registry_path = PROFILING_DIR / "01_meter_registry.csv"
    stats_path = PROFILING_DIR / "07_meter_statistics.csv"
    
    if not registry_path.exists() or not stats_path.exists():
        print("필요한 데이터 파일이 없습니다. meter_detailed_analysis.py를 먼저 실행하세요.")
        return
        
    df_registry = pd.read_csv(registry_path).sort_values(["equipment_group", "meter_urn"])
    df_stats = pd.read_csv(stats_path)
    
    # Word 문서 객체 생성
    doc = Document()
    
    # 여백 설정 (전체 0.5인치로 가독성 및 공간 확보)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # 표지
    doc.add_heading("81개 계량기 상세 분석 및 요약 리포트", 0)
    p = doc.add_paragraph()
    p.add_run("Honda R&D Europe 건물에 설치된 81개 개별 계량기 각각의 데이터 기반 자동 분석 리포트입니다.\n")
    p.add_run("각 페이지는 하나의 계량기 정보를 담고 있으며, 시계열 패턴과 통계량 기반의 해석(Insight)을 제공합니다.")
    doc.add_page_break()
    
    # 계량기별 페이지 생성
    for idx, row in df_registry.iterrows():
        urn = row["meter_urn"]
        group = row["equipment_group"]
        
        doc.add_heading(f"계량기: {urn} ({group})", level=1)
        
        meter_stats = df_stats[df_stats["meter_urn"] == urn]
        
        # 인사이트 (해석) 텍스트 추가
        doc.add_heading("데이터 분석 요약 (Insight)", level=3)
        if not meter_stats.empty:
            insights = analyze_meter(meter_stats, group)
            for insight in insights:
                doc.add_paragraph(insight, style="List Bullet")
        else:
            doc.add_paragraph("통계 데이터가 충분하지 않아 분석할 수 없습니다.")
            
        doc.add_paragraph() # 여백
            
        # 통계 표 추가
        doc.add_heading("기초 통계량 (1시간 단위)", level=3)
        if not meter_stats.empty:
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "측정항목"
            hdr_cells[1].text = "데이터 수"
            hdr_cells[2].text = "평균"
            hdr_cells[3].text = "최소"
            hdr_cells[4].text = "최대"
            hdr_cells[5].text = "무부하 비율"
            
            for _, stat in meter_stats.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(stat['measurement'])
                row_cells[1].text = f"{int(stat['count']):,}"
                row_cells[2].text = f"{stat['mean']:.2f}" if pd.notnull(stat['mean']) else "-"
                row_cells[3].text = f"{stat['min']:.2f}" if pd.notnull(stat['min']) else "-"
                row_cells[4].text = f"{stat['max']:.2f}" if pd.notnull(stat['max']) else "-"
                zero_pct = f"{stat['zero_ratio']*100:.1f}%" if pd.notnull(stat['zero_ratio']) else "-"
                row_cells[5].text = zero_pct
        else:
            doc.add_paragraph("통계 데이터가 없습니다.")
            
        doc.add_paragraph() # 여백
            
        # 차트 이미지 삽입
        chart_filename = f"{urn}_weekly.png"
        chart_local_path = FIGURES_DIR / chart_filename
        
        doc.add_heading("주간(Weekly) 트렌드 차트", level=3)
        if chart_local_path.exists():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(str(chart_local_path), width=Inches(6.5))
        else:
            doc.add_paragraph("(시각화 차트가 존재하지 않습니다.)")
            
        # 다음 계량기는 새 페이지에서 시작 (마지막 계량기 제외)
        if idx != len(df_registry) - 1:
            doc.add_page_break()

    out_docx = DOCS_DIR / "07_81개_계량기_개별_분석_개선본.docx"
    doc.save(out_docx)
    print(f"▶ 레이아웃 최적화 리포트 생성 완료: {out_docx}")


if __name__ == "__main__":
    generate_docx()
