"""
일일 보고서 문서 변환 모듈 — PDF / DOCX / HWPX.

- PDF : reportlab (Docker 이미지에 fonts-nanum 설치 → 한글 정상 출력)
- DOCX: python-docx (한글에서도 열림, 표 포함)
- HWPX: OWPML(ZIP+XML) 직접 조립. 한글 2014+ 에서 열림.
        단순 문단 구조만 사용 (표 미사용) — 서식 단순화로 호환성 확보.

각 build_*(report: dict) → bytes 를 반환한다.
report dict 는 report.py 의 _fetch_daily / build_daily_report 결과 형식을 따른다.
"""

import io
import zipfile
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape as _xml_escape

_NANUM_PATHS = [
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
    "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf",
    "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
]


# ── 공통: 보고서 → 표/요약 데이터 ────────────────────────────────

def _kpi_rows(r: dict) -> list[tuple[str, str]]:
    peak = (
        f"{r['peak_hour']}시 ({r['peak_kw']:,.0f} kW)"
        if r.get("peak_hour") is not None else "—"
    )
    def _num(v, suffix="", fmt="{:,.1f}"):
        return (fmt.format(v) + suffix) if v is not None else "—"

    return [
        ("총 소비량",     _num(r.get("total_consumption_kwh"), " kWh", "{:,.0f}")),
        ("계통 전력 비용", _num(r.get("cost_eur"), "", "€ {:,.0f}")),
        ("CO₂ 배출",      _num(r.get("co2_kg"), " kg", "{:,.0f}")),
        ("자급률",        _num(r.get("self_sufficiency_pct"), "%")),
        ("평균 COP",      _num(r.get("avg_cop"), "", "{:.2f}")),
        ("그리드 의존도", _num(r.get("grid_dependency_pct"), "%")),
        ("PV 발전",       _num(r.get("pv_kwh"), " kWh", "{:,.0f}")),
        ("CHP 발전",      _num(r.get("chp_kwh"), " kWh", "{:,.0f}")),
        ("피크 시간",     peak),
        ("이상탐지",      f"{r.get('anomaly_count', 0)}건"),
    ]


def _hourly_rows(r: dict) -> list[list[str]]:
    rows = []
    for h in (r.get("hourly_profile") or []):
        def _c(v):
            return f"{v:,.1f}" if v is not None else "—"
        rows.append([
            f"{h.get('hour', 0):02d}시",
            _c(h.get("grid_kw")), _c(h.get("pv_kw")), _c(h.get("chp_kw")),
            _c(h.get("total_kw")),
            f"{h['cop']:.2f}" if h.get("cop") is not None else "—",
        ])
    return rows


_HOURLY_HEADER = ["시간", "계통(kW)", "태양광(kW)", "CHP(kW)", "합계(kW)", "COP"]

_TYPE_LABEL = {
    "COPDrop": "COP 급락", "CHPOutage": "CHP 정지", "PowerSpike": "전력 급등",
    "NightConsumption": "야간 소비", "PVNightNonZero": "PV 야간 비정상",
}
_EVENT_HEADER = ["시각", "유형", "심각도", "계측기", "설명"]


def _event_rows(r: dict) -> list[list[str]]:
    rows = []
    for e in (r.get("anomaly_events") or []):
        ts = (e.get("timestamp") or "")[11:16]  # HH:MM
        rows.append([
            ts,
            _TYPE_LABEL.get(e.get("anomaly_type"), e.get("anomaly_type") or "—"),
            e.get("severity") or "—",
            e.get("meter_id") or "—",
            (e.get("description") or "—")[:60],
        ])
    return rows


# ── 차트 (matplotlib, Agg) — 라벨은 폰트 안전하게 영문 ───────────

def _mpl_plt():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    return plt


def _chart_hourly_png(r: dict) -> bytes | None:
    """시간대별 전력 프로파일 (Grid/PV/CHP 스택)."""
    hourly = r.get("hourly_profile") or []
    if not hourly:
        return None
    try:
        plt = _mpl_plt()
        hours = [h.get("hour", i) for i, h in enumerate(hourly)]
        grid  = [h.get("grid_kw") or 0 for h in hourly]
        pv    = [h.get("pv_kw") or 0 for h in hourly]
        chp   = [h.get("chp_kw") or 0 for h in hourly]
        fig, ax = plt.subplots(figsize=(7.4, 2.5), dpi=130)
        ax.stackplot(hours, grid, pv, chp, labels=["Grid", "PV", "CHP"],
                     colors=["#1f6feb", "#f0b429", "#16a34a"], alpha=0.9)
        ax.set_xlabel("Hour", fontsize=8); ax.set_ylabel("kW", fontsize=8)
        ax.set_xlim(0, 23); ax.set_xticks(range(0, 24, 2))
        ax.tick_params(labelsize=7)
        ax.legend(loc="upper left", fontsize=8, ncol=3, framealpha=0.6)
        ax.grid(True, alpha=0.2)
        fig.tight_layout()
        buf = io.BytesIO(); fig.savefig(buf, format="png"); plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None


def _chart_mix_png(r: dict) -> bytes | None:
    """에너지 믹스 도넛 (Grid/PV/CHP kWh)."""
    pv = r.get("pv_kwh") or 0
    chp = r.get("chp_kwh") or 0
    grid = r.get("grid_kwh")
    if grid is None:
        total = r.get("total_consumption_kwh") or 0
        grid = max(total - pv - chp, 0)
    vals = [grid, pv, chp]
    if sum(vals) <= 0:
        return None
    try:
        plt = _mpl_plt()
        fig, ax = plt.subplots(figsize=(3.0, 2.5), dpi=130)
        ax.pie(vals, labels=["Grid", "PV", "CHP"],
               colors=["#1f6feb", "#f0b429", "#16a34a"],
               autopct="%1.0f%%", startangle=90,
               wedgeprops={"width": 0.42}, textprops={"fontsize": 8})
        ax.set_title("Energy Mix", fontsize=9)
        fig.tight_layout()
        buf = io.BytesIO(); fig.savefig(buf, format="png"); plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None


# ── PDF (reportlab) ──────────────────────────────────────────────

def build_pdf(r: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )

    font = "Helvetica"
    for p in _NANUM_PATHS:
        if Path(p).exists():
            try:
                pdfmetrics.registerFont(TTFont("Kor", p))
                font = "Kor"
            except Exception:
                pass
            break

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle("t", parent=styles["Title"],    fontName=font, fontSize=16)
    body_s  = ParagraphStyle("b", parent=styles["Normal"],   fontName=font, fontSize=10, leading=16)
    head_s  = ParagraphStyle("h", parent=styles["Heading2"], fontName=font, fontSize=12)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18*mm, rightMargin=18*mm,
                            topMargin=18*mm, bottomMargin=18*mm)
    story = [
        Paragraph("일일 에너지 보고서", title_s),
        Paragraph(f"Honda R&D Europe GmbH — {r.get('date', '')}", body_s),
        Paragraph(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_s),
        Spacer(1, 5*mm),
        Paragraph("KPI 요약", head_s),
    ]

    kpi_tbl = Table([["항목", "값"]] + [list(t) for t in _kpi_rows(r)], colWidths=[60*mm, 100*mm])
    kpi_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f6feb")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTNAME",   (0, 0), (-1, -1), font),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f8fa")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d7de")),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story += [kpi_tbl, Spacer(1, 5*mm)]

    # 차트 — 시간대별 프로파일 + 에너지 믹스
    from reportlab.platypus import Image as _RLImage
    _hourly_png = _chart_hourly_png(r)
    _mix_png    = _chart_mix_png(r)
    if _hourly_png or _mix_png:
        story += [Paragraph("그래프", head_s)]
        if _hourly_png:
            story += [_RLImage(io.BytesIO(_hourly_png), width=170*mm, height=57*mm), Spacer(1, 3*mm)]
        if _mix_png:
            story += [_RLImage(io.BytesIO(_mix_png), width=70*mm, height=58*mm)]
        story.append(Spacer(1, 5*mm))

    if r.get("ai_summary"):
        story += [Paragraph("AI 요약", head_s)]
        for line in r["ai_summary"].split("\n"):
            story.append(Paragraph(line.strip() or "&nbsp;", body_s))
        story.append(Spacer(1, 5*mm))

    events = _event_rows(r)
    story += [Paragraph(f"당일 이상 이벤트 ({len(events)}건)", head_s)]
    if events:
        e_tbl = Table([_EVENT_HEADER] + events, repeatRows=1,
                      colWidths=[18*mm, 26*mm, 20*mm, 28*mm, 68*mm])
        e_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#8b1a1a")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME",   (0, 0), (-1, -1), font),
            ("FONTSIZE",   (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fbf0f0")]),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d0d7de")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 2), ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        story.append(e_tbl)
    else:
        story.append(Paragraph("탐지된 이상 없음", body_s))
    story.append(Spacer(1, 5*mm))

    hourly = _hourly_rows(r)
    if hourly:
        story += [Paragraph("시간대별 전력 프로파일", head_s)]
        h_tbl = Table([_HOURLY_HEADER] + hourly, repeatRows=1)
        h_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#30363d")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME",   (0, 0), (-1, -1), font),
            ("FONTSIZE",   (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f8fa")]),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d0d7de")),
            ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
            ("TOPPADDING", (0, 0), (-1, -1), 2), ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        story.append(h_tbl)

    doc.build(story)
    return buf.getvalue()


# ── DOCX (python-docx) ───────────────────────────────────────────

def build_docx(r: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()
    doc.add_heading("일일 에너지 보고서", level=0)
    doc.add_paragraph(f"Honda R&D Europe GmbH — {r.get('date', '')}")
    doc.add_paragraph(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("KPI 요약", level=2)
    kpi = _kpi_rows(r)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text, t.rows[0].cells[1].text = "항목", "값"
    for label, value in kpi:
        c = t.add_row().cells
        c[0].text, c[1].text = label, value

    # 차트
    _hourly_png = _chart_hourly_png(r)
    _mix_png    = _chart_mix_png(r)
    if _hourly_png or _mix_png:
        doc.add_heading("그래프", level=2)
        if _hourly_png:
            doc.add_picture(io.BytesIO(_hourly_png), width=Inches(6.3))
        if _mix_png:
            doc.add_picture(io.BytesIO(_mix_png), width=Inches(2.6))

    if r.get("ai_summary"):
        doc.add_heading("AI 요약", level=2)
        doc.add_paragraph(r["ai_summary"])

    events = _event_rows(r)
    doc.add_heading(f"당일 이상 이벤트 ({len(events)}건)", level=2)
    if events:
        et = doc.add_table(rows=1, cols=len(_EVENT_HEADER))
        et.style = "Light Grid Accent 2"
        for i, h in enumerate(_EVENT_HEADER):
            et.rows[0].cells[i].text = h
        for row in events:
            cells = et.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = v
    else:
        doc.add_paragraph("탐지된 이상 없음")

    hourly = _hourly_rows(r)
    if hourly:
        doc.add_heading("시간대별 전력 프로파일", level=2)
        ht = doc.add_table(rows=1, cols=len(_HOURLY_HEADER))
        ht.style = "Light Grid Accent 1"
        for i, h in enumerate(_HOURLY_HEADER):
            ht.rows[0].cells[i].text = h
        for row in hourly:
            cells = ht.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = v

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HWPX (OWPML ZIP) ─────────────────────────────────────────────

def _hx(s) -> str:
    return _xml_escape(str(s))


_VERSION_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<hv:HCFVersion xmlns:hv="http://www.hancom.co.kr/hwpml/2011/version" tagetApplication="WORDPROCESSOR" major="5" minor="0" micro="5" buildNumber="0" os="1" xmlVersion="1.4" application="EMS Agent" appVersion="1.0"/>"""

_CONTAINER_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<ocf:container xmlns:ocf="urn:oasis:names:tc:opendocument:xmlns:container">
  <ocf:rootfiles>
    <ocf:rootfile full-path="Contents/content.hpf" media-type="application/hwpml-package+xml"/>
  </ocf:rootfiles>
</ocf:container>"""

_MANIFEST_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<odf:manifest xmlns:odf="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0">
  <odf:file-entry full-path="Contents/content.hpf" media-type="application/hwpml-package+xml"/>
  <odf:file-entry full-path="Contents/header.xml" media-type="application/xml"/>
  <odf:file-entry full-path="Contents/section0.xml" media-type="application/xml"/>
  <odf:file-entry full-path="settings.xml" media-type="application/xml"/>
</odf:manifest>"""

_SETTINGS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<ha:HWPApplicationSetting xmlns:ha="http://www.hancom.co.kr/hwpml/2011/app"/>"""

_CONTENT_HPF = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<hpf:package xmlns:hpf="http://www.hancom.co.kr/schema/2011/hpf" xmlns:opf="http://www.idpf.org/2007/opf/" version="">
  <hpf:metadata>
    <opf:title>일일 에너지 보고서</opf:title>
  </hpf:metadata>
  <hpf:manifest>
    <hpf:item id="header" href="Contents/header.xml" media-type="application/xml"/>
    <hpf:item id="section0" href="Contents/section0.xml" media-type="application/xml"/>
    <hpf:item id="settings" href="settings.xml" media-type="application/xml"/>
  </hpf:manifest>
  <hpf:spine>
    <hpf:itemref idref="header"/>
    <hpf:itemref idref="section0"/>
  </hpf:spine>
</hpf:package>"""

# 7개 언어 폰트 매핑 (HANGUL / LATIN / HANJA / JAPANESE / OTHER / SYMBOL / USER)
_FONT_LANGS = ["HANGUL", "LATIN", "HANJA", "JAPANESE", "OTHER", "SYMBOL", "USER"]


def _header_xml() -> str:
    fontfaces = []
    for lang in _FONT_LANGS:
        fontfaces.append(
            f'<hh:fontface lang="{lang}" fontCnt="1">'
            f'<hh:font id="0" face="함초롬바탕" type="TTF" isEmbedded="0">'
            f'<hh:typeInfo familyType="FCAT_GOTHIC" weight="0" proportion="0" contrast="0" '
            f'strokeVariation="0" armStyle="0" letterform="0" midline="0" xHeight="0"/>'
            f'</hh:font></hh:fontface>'
        )
    # charPr: 7개 언어 모두 fontRef=0
    lang_refs = "".join(f'<hh:{tag} hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>'
                        for tag in ["fontRef", "ratio", "spacing", "relSz", "offset"])
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" version="1.4" secCnt="1">
<hh:beginNum page="1" footnote="1" endnote="1" pic="1" tbl="1" equation="1"/>
<hh:refList>
<hh:fontfaces itemCnt="7">{''.join(fontfaces)}</hh:fontfaces>
<hh:borderFills itemCnt="1">
<hh:borderFill id="1" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0">
<hh:slash type="NONE" Crooked="0" isCounter="0"/><hh:backSlash type="NONE" Crooked="0" isCounter="0"/>
<hh:leftBorder type="NONE" width="0.1 mm" color="#000000"/><hh:rightBorder type="NONE" width="0.1 mm" color="#000000"/>
<hh:topBorder type="NONE" width="0.1 mm" color="#000000"/><hh:bottomBorder type="NONE" width="0.1 mm" color="#000000"/>
<hh:diagonal type="SOLID" width="0.1 mm" color="#000000"/>
</hh:borderFill>
</hh:borderFills>
<hh:charProperties itemCnt="1">
<hh:charPr id="0" height="1000" textColor="#000000" shadeColor="none" useFontSpace="0" useKerning="0" symMark="NONE" borderFillIDRef="1">
{lang_refs}
<hh:bold/><hh:underline type="NONE" shape="SOLID" color="#000000"/>
<hh:strikeout shape="NONE" color="#000000"/><hh:outline type="NONE"/><hh:shadow type="NONE" color="#C0C0C0" offsetX="10" offsetY="10"/>
</hh:charPr>
</hh:charProperties>
<hh:tabProperties itemCnt="1">
<hh:tabPr id="0" autoTabLeft="0" autoTabRight="0"/>
</hh:tabProperties>
<hh:numberings itemCnt="1">
<hh:numbering id="1" start="0"><hh:paraHead start="1" level="1" align="LEFT" useInstWidth="1" autoIndent="1" widthAdjust="0" textOffsetType="PERCENT" textOffset="50" numFormat="DIGIT" charPrIDRef="4294967295">^1.</hh:paraHead></hh:numbering>
</hh:numberings>
<hh:paraProperties itemCnt="1">
<hh:paraPr id="0" tabPrIDRef="0" condense="0" fontLineHeight="0" snapToGrid="1" suppressLineNumbers="0" checked="0">
<hh:align horizontal="LEFT" vertical="BASELINE"/>
<hh:heading type="NONE" idRef="0" level="0"/>
<hh:breakSetting breakLatinWord="KEEP_WORD" breakNonLatinWord="KEEP_WORD" widowOrphan="0" keepWithNext="0" keepLines="0" pageBreakBefore="0" lineWrap="BREAK"/>
<hh:margin><hc:intent value="0" unit="HWPUNIT"/><hc:left value="0" unit="HWPUNIT"/><hc:right value="0" unit="HWPUNIT"/><hc:prev value="0" unit="HWPUNIT"/><hc:next value="0" unit="HWPUNIT"/></hh:margin>
<hh:lineSpacing type="PERCENT" value="160" unit="HWPUNIT"/>
<hh:border borderFillIDRef="1" offsetLeft="0" offsetRight="0" offsetTop="0" offsetBottom="0" connect="0" ignoreMargin="0"/>
</hh:paraPr>
</hh:paraProperties>
<hh:styles itemCnt="1">
<hh:style id="0" type="PARA" name="바탕글" engName="Normal" paraPrIDRef="0" charPrIDRef="0" nextStyleIDRef="0" langID="1042" lockForm="0"/>
</hh:styles>
</hh:refList>
</hh:head>"""


def _para(text: str, bold: bool = False) -> str:
    """단순 문단 한 개 생성. (charPr 0 사용, 굵게는 별도 미지원 — 단순화)"""
    return (
        f'<hp:p paraPrIDRef="0" styleIDRef="0">'
        f'<hp:run charPrIDRef="0"><hp:t>{_hx(text)}</hp:t></hp:run>'
        f'</hp:p>'
    )


def _section_xml(r: dict) -> str:
    paras = []
    # 첫 문단에 구역 정의(secPr) 포함
    first = (
        '<hp:p paraPrIDRef="0" styleIDRef="0"><hp:run charPrIDRef="0">'
        '<hp:secPr id="0" textDirection="HORIZONTAL" spaceColumns="0" tabStop="8000" tabStopVal="4000" tabStopUnit="HWPUNIT" outlineShapeIDRef="0" memoShapeIDRef="0" textVerticalWidthHead="0">'
        '<hp:grid lineGrid="0" charGrid="0" wonggojiFormat="0" strtnum="1"/>'
        '<hp:startNum pageStartsOn="BOTH" page="0" pic="0" tbl="0" equation="0"/>'
        '<hp:pagePr landscape="WIDELY" width="59528" height="84188" gutterType="LEFT_ONLY">'
        '<hp:margin header="4252" footer="4252" gutter="0" left="8504" right="8504" top="5668" bottom="4252"/>'
        '</hp:pagePr>'
        '</hp:secPr>'
        f'<hp:t>{_hx("일일 에너지 보고서")}</hp:t></hp:run></hp:p>'
    )
    paras.append(first)
    paras.append(_para(f"Honda R&D Europe GmbH — {r.get('date', '')}"))
    paras.append(_para(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}"))
    paras.append(_para(""))
    paras.append(_para("[ KPI 요약 ]"))
    for label, value in _kpi_rows(r):
        paras.append(_para(f"  · {label}: {value}"))

    if r.get("ai_summary"):
        paras.append(_para(""))
        paras.append(_para("[ AI 요약 ]"))
        for line in r["ai_summary"].split("\n"):
            if line.strip():
                paras.append(_para("  " + line.strip()))

    events = _event_rows(r)
    paras.append(_para(""))
    paras.append(_para(f"[ 당일 이상 이벤트 ({len(events)}건) ]"))
    if events:
        paras.append(_para("  " + " | ".join(_EVENT_HEADER)))
        for row in events:
            paras.append(_para("  " + " | ".join(row)))
    else:
        paras.append(_para("  탐지된 이상 없음"))

    hourly = _hourly_rows(r)
    if hourly:
        paras.append(_para(""))
        paras.append(_para("[ 시간대별 전력 프로파일 ]"))
        paras.append(_para("  " + " | ".join(_HOURLY_HEADER)))
        for row in hourly:
            paras.append(_para("  " + " | ".join(row)))

    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" '
        'xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph" '
        'xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">'
        + "".join(paras) +
        '</hs:sec>'
    )


def build_hwpx(r: dict) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        # mimetype 은 반드시 첫 엔트리 + 무압축
        zi = zipfile.ZipInfo("mimetype")
        zi.compress_type = zipfile.ZIP_STORED
        z.writestr(zi, "application/hwp+zip")

        z.writestr("version.xml", _VERSION_XML)
        z.writestr("settings.xml", _SETTINGS_XML)
        z.writestr("Contents/content.hpf", _CONTENT_HPF)
        z.writestr("Contents/header.xml", _header_xml())
        z.writestr("Contents/section0.xml", _section_xml(r))
        z.writestr("META-INF/container.xml", _CONTAINER_XML)
        z.writestr("META-INF/manifest.xml", _MANIFEST_XML)
    return buf.getvalue()


# ── 디스패처 ─────────────────────────────────────────────────────

_BUILDERS = {"pdf": build_pdf, "docx": build_docx, "hwpx": build_hwpx}
_MEDIA = {
    "pdf":  "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "hwpx": "application/hwp+zip",
}


def render(report: dict, fmt: str) -> tuple[bytes, str, str]:
    """(bytes, media_type, filename) 반환. 지원하지 않는 포맷이면 ValueError."""
    fmt = fmt.lower()
    if fmt not in _BUILDERS:
        raise ValueError(f"지원하지 않는 포맷: {fmt}")
    data = _BUILDERS[fmt](report)
    filename = f"daily_report_{report.get('date', 'report')}.{fmt}"
    return data, _MEDIA[fmt], filename


# ══════════════════════════════════════════════════════════════════
#  월간 보고서 (Monthly Report) — PDF / DOCX
# ══════════════════════════════════════════════════════════════════

_M_ENERGY_PRICE = 0.20
_M_GRID_CO2     = 0.38
_M_CHP_CO2      = 0.20


def _monthly_cost_co2(it: dict) -> tuple[float, float]:
    """월 KPI에서 비용(€)·CO₂(kg) 파생."""
    total = it.get("total_consumption_kwh") or 0
    grid_dep = it.get("grid_dependency_pct")
    chp = it.get("chp_kwh") or 0
    grid_kwh = total * grid_dep / 100.0 if grid_dep is not None else 0
    return grid_kwh * _M_ENERGY_PRICE, grid_kwh * _M_GRID_CO2 + chp * _M_CHP_CO2


_MONTHLY_HEADER = ["월", "소비(kWh)", "비용(€)", "CO₂(kg)", "자급률", "COP", "그리드", "이상"]


def _monthly_rows(items: list[dict]) -> list[list[str]]:
    rows = []
    for it in items:
        cost, co2 = _monthly_cost_co2(it)
        def _n(v, fmt="{:,.0f}", suf=""):
            return (fmt.format(v) + suf) if v is not None else "—"
        rows.append([
            it.get("period", "—"),
            _n(it.get("total_consumption_kwh")),
            _n(round(cost)),
            _n(round(co2)),
            _n(it.get("self_sufficiency_pct"), "{:.1f}", "%"),
            _n(it.get("avg_cop"), "{:.2f}"),
            _n(it.get("grid_dependency_pct"), "{:.1f}", "%"),
            f"{it.get('anomaly_count', 0)}건",
        ])
    return rows


def _chart_monthly_trend_png(items: list[dict]) -> bytes | None:
    """월별 소비량(막대) + 자급률(선) 추이."""
    if not items:
        return None
    try:
        plt = _mpl_plt()
        periods = [it.get("period", "") for it in items]
        cons    = [it.get("total_consumption_kwh") or 0 for it in items]
        ss      = [it.get("self_sufficiency_pct") for it in items]
        fig, ax1 = plt.subplots(figsize=(7.4, 2.8), dpi=130)
        x = range(len(periods))
        ax1.bar(x, cons, color="#1f6feb", alpha=0.8, label="Consumption (kWh)")
        ax1.set_ylabel("kWh", fontsize=8); ax1.tick_params(labelsize=7)
        ax1.set_xticks(list(x)); ax1.set_xticklabels(periods, rotation=45, ha="right", fontsize=6)
        ax2 = ax1.twinx()
        ax2.plot(x, [v if v is not None else None for v in ss], color="#16a34a",
                 marker="o", markersize=3, linewidth=1.8, label="Self-sufficiency (%)")
        ax2.set_ylabel("%", fontsize=8); ax2.tick_params(labelsize=7)
        ax2.set_ylim(0, max([v for v in ss if v is not None] + [50]) * 1.3)
        l1, lb1 = ax1.get_legend_handles_labels()
        l2, lb2 = ax2.get_legend_handles_labels()
        ax1.legend(l1 + l2, lb1 + lb2, loc="upper left", fontsize=7, framealpha=0.6)
        ax1.grid(True, axis="y", alpha=0.2)
        fig.tight_layout()
        buf = io.BytesIO(); fig.savefig(buf, format="png"); plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None


def _kor_font_pdf():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    for p in _NANUM_PATHS:
        if Path(p).exists():
            try:
                pdfmetrics.registerFont(TTFont("Kor", p)); return "Kor"
            except Exception:
                pass
            break
    return "Helvetica"


def build_monthly_pdf(payload: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage,
    )

    items = payload.get("items", [])
    font = _kor_font_pdf()
    styles = getSampleStyleSheet()
    title_s = ParagraphStyle("t", parent=styles["Title"],    fontName=font, fontSize=16)
    body_s  = ParagraphStyle("b", parent=styles["Normal"],   fontName=font, fontSize=10, leading=16)
    head_s  = ParagraphStyle("h", parent=styles["Heading2"], fontName=font, fontSize=12)

    period_range = f"{items[0]['period']} ~ {items[-1]['period']}" if items else ""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18*mm, rightMargin=18*mm,
                            topMargin=18*mm, bottomMargin=18*mm)
    story = [
        Paragraph("월간 에너지 보고서", title_s),
        Paragraph(f"Honda R&D Europe GmbH — {period_range}", body_s),
        Paragraph(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_s),
        Spacer(1, 5*mm),
    ]

    trend = _chart_monthly_trend_png(items)
    if trend:
        story += [Paragraph("월별 추이", head_s),
                  RLImage(io.BytesIO(trend), width=170*mm, height=64*mm), Spacer(1, 5*mm)]

    story += [Paragraph("월별 KPI", head_s)]
    tbl = Table([_MONTHLY_HEADER] + _monthly_rows(items), repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f6feb")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTNAME",   (0, 0), (-1, -1), font),
        ("FONTSIZE",   (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f8fa")]),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d0d7de")),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story += [tbl, Spacer(1, 5*mm)]

    if payload.get("trend_narrative"):
        story += [Paragraph("AI 트렌드 분석", head_s)]
        for line in payload["trend_narrative"].split("\n"):
            story.append(Paragraph(line.strip() or "&nbsp;", body_s))

    doc.build(story)
    return buf.getvalue()


def build_monthly_docx(payload: dict) -> bytes:
    from docx import Document
    from docx.shared import Inches

    items = payload.get("items", [])
    period_range = f"{items[0]['period']} ~ {items[-1]['period']}" if items else ""
    doc = Document()
    doc.add_heading("월간 에너지 보고서", level=0)
    doc.add_paragraph(f"Honda R&D Europe GmbH — {period_range}")
    doc.add_paragraph(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    trend = _chart_monthly_trend_png(items)
    if trend:
        doc.add_heading("월별 추이", level=2)
        doc.add_picture(io.BytesIO(trend), width=Inches(6.3))

    doc.add_heading("월별 KPI", level=2)
    t = doc.add_table(rows=1, cols=len(_MONTHLY_HEADER))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(_MONTHLY_HEADER):
        t.rows[0].cells[i].text = h
    for row in _monthly_rows(items):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    if payload.get("trend_narrative"):
        doc.add_heading("AI 트렌드 분석", level=2)
        doc.add_paragraph(payload["trend_narrative"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_MONTHLY_BUILDERS = {"pdf": build_monthly_pdf, "docx": build_monthly_docx}


def render_monthly(payload: dict, fmt: str) -> tuple[bytes, str, str]:
    fmt = fmt.lower()
    if fmt not in _MONTHLY_BUILDERS:
        raise ValueError(f"월간 보고서 미지원 포맷: {fmt} (pdf/docx만)")
    data = _MONTHLY_BUILDERS[fmt](payload)
    items = payload.get("items", [])
    tag = items[-1]["period"] if items else "report"
    return data, _MEDIA[fmt], f"monthly_report_{tag}.{fmt}"
