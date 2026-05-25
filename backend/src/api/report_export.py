"""
일일 보고서 문서 변환 모듈 — PDF / DOCX / HWPX.

- PDF : reportlab (Docker 이미지에 fonts-nanum 설치 → 한글 정상 출력)
- DOCX: python-docx (한글에서도 열림, 표 포함)
- HWPX: OWPML(ZIP+XML) 직접 조립. 한글 2014+ 에서 열림.
        단순 문단 구조만 사용 (표 미사용) — 서식 단순화로 호환성 확보.

각 build_*(report: dict) → bytes 를 반환한다.
report dict 는 report.py 의 _fetch_daily / build_daily_report 결과 형식을 따른다.
"""

import io
import zipfile
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape as _xml_escape

_NANUM_PATHS = [
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
    "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf",
    "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
]


# ── 공통: 보고서 → 표/요약 데이터 ────────────────────────────────

def _kpi_rows(r: dict) -> list[tuple[str, str]]:
    peak = (
        f"{r['peak_hour']}시 ({r['peak_kw']:,.0f} kW)"
        if r.get("peak_hour") is not None else "—"
    )
    def _num(v, suffix="", fmt="{:,.1f}"):
        return (fmt.format(v) + suffix) if v is not None else "—"

    return [
        ("총 소비량",     _num(r.get("total_consumption_kwh"), " kWh", "{:,.0f}")),
        ("자급률",        _num(r.get("self_sufficiency_pct"), "%")),
        ("평균 COP",      _num(r.get("avg_cop"), "", "{:.2f}")),
        ("그리드 의존도", _num(r.get("grid_dependency_pct"), "%")),
        ("PV 발전",       _num(r.get("pv_kwh"), " kWh", "{:,.0f}")),
        ("CHP 발전",      _num(r.get("chp_kwh"), " kWh", "{:,.0f}")),
        ("피크 시간",     peak),
        ("이상탐지",      f"{r.get('anomaly_count', 0)}건"),
    ]


def _hourly_rows(r: dict) -> list[list[str]]:
    rows = []
    for h in (r.get("hourly_profile") or []):
        def _c(v):
            return f"{v:,.1f}" if v is not None else "—"
        rows.append([
            f"{h.get('hour', 0):02d}시",
            _c(h.get("grid_kw")), _c(h.get("pv_kw")), _c(h.get("chp_kw")),
            _c(h.get("total_kw")),
            f"{h['cop']:.2f}" if h.get("cop") is not None else "—",
        ])
    return rows


_HOURLY_HEADER = ["시간", "계통(kW)", "태양광(kW)", "CHP(kW)", "합계(kW)", "COP"]


# ── PDF (reportlab) ──────────────────────────────────────────────

def build_pdf(r: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )

    font = "Helvetica"
    for p in _NANUM_PATHS:
        if Path(p).exists():
            try:
                pdfmetrics.registerFont(TTFont("Kor", p))
                font = "Kor"
            except Exception:
                pass
            break

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle("t", parent=styles["Title"],    fontName=font, fontSize=16)
    body_s  = ParagraphStyle("b", parent=styles["Normal"],   fontName=font, fontSize=10, leading=16)
    head_s  = ParagraphStyle("h", parent=styles["Heading2"], fontName=font, fontSize=12)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18*mm, rightMargin=18*mm,
                            topMargin=18*mm, bottomMargin=18*mm)
    story = [
        Paragraph("일일 에너지 보고서", title_s),
        Paragraph(f"Honda R&D Europe GmbH — {r.get('date', '')}", body_s),
        Paragraph(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_s),
        Spacer(1, 5*mm),
        Paragraph("KPI 요약", head_s),
    ]

    kpi_tbl = Table([["항목", "값"]] + [list(t) for t in _kpi_rows(r)], colWidths=[60*mm, 100*mm])
    kpi_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f6feb")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTNAME",   (0, 0), (-1, -1), font),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f8fa")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d7de")),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story += [kpi_tbl, Spacer(1, 5*mm)]

    if r.get("ai_summary"):
        story += [Paragraph("AI 요약", head_s)]
        for line in r["ai_summary"].split("\n"):
            story.append(Paragraph(line.strip() or "&nbsp;", body_s))
        story.append(Spacer(1, 5*mm))

    hourly = _hourly_rows(r)
    if hourly:
        story += [Paragraph("시간대별 전력 프로파일", head_s)]
        h_tbl = Table([_HOURLY_HEADER] + hourly, repeatRows=1)
        h_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#30363d")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTNAME",   (0, 0), (-1, -1), font),
            ("FONTSIZE",   (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f8fa")]),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d0d7de")),
            ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
            ("TOPPADDING", (0, 0), (-1, -1), 2), ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        story.append(h_tbl)

    doc.build(story)
    return buf.getvalue()


# ── DOCX (python-docx) ───────────────────────────────────────────

def build_docx(r: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("일일 에너지 보고서", level=0)
    doc.add_paragraph(f"Honda R&D Europe GmbH — {r.get('date', '')}")
    doc.add_paragraph(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("KPI 요약", level=2)
    kpi = _kpi_rows(r)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text, t.rows[0].cells[1].text = "항목", "값"
    for label, value in kpi:
        c = t.add_row().cells
        c[0].text, c[1].text = label, value

    if r.get("ai_summary"):
        doc.add_heading("AI 요약", level=2)
        doc.add_paragraph(r["ai_summary"])

    hourly = _hourly_rows(r)
    if hourly:
        doc.add_heading("시간대별 전력 프로파일", level=2)
        ht = doc.add_table(rows=1, cols=len(_HOURLY_HEADER))
        ht.style = "Light Grid Accent 1"
        for i, h in enumerate(_HOURLY_HEADER):
            ht.rows[0].cells[i].text = h
        for row in hourly:
            cells = ht.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = v

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HWPX (OWPML ZIP) ─────────────────────────────────────────────

def _hx(s) -> str:
    return _xml_escape(str(s))


_VERSION_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<hv:HCFVersion xmlns:hv="http://www.hancom.co.kr/hwpml/2011/version" tagetApplication="WORDPROCESSOR" major="5" minor="0" micro="5" buildNumber="0" os="1" xmlVersion="1.4" application="EMS Agent" appVersion="1.0"/>"""

_CONTAINER_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<ocf:container xmlns:ocf="urn:oasis:names:tc:opendocument:xmlns:container">
  <ocf:rootfiles>
    <ocf:rootfile full-path="Contents/content.hpf" media-type="application/hwpml-package+xml"/>
  </ocf:rootfiles>
</ocf:container>"""

_MANIFEST_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<odf:manifest xmlns:odf="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0">
  <odf:file-entry full-path="Contents/content.hpf" media-type="application/hwpml-package+xml"/>
  <odf:file-entry full-path="Contents/header.xml" media-type="application/xml"/>
  <odf:file-entry full-path="Contents/section0.xml" media-type="application/xml"/>
  <odf:file-entry full-path="settings.xml" media-type="application/xml"/>
</odf:manifest>"""

_SETTINGS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<ha:HWPApplicationSetting xmlns:ha="http://www.hancom.co.kr/hwpml/2011/app"/>"""

_CONTENT_HPF = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<hpf:package xmlns:hpf="http://www.hancom.co.kr/schema/2011/hpf" xmlns:opf="http://www.idpf.org/2007/opf/" version="">
  <hpf:metadata>
    <opf:title>일일 에너지 보고서</opf:title>
  </hpf:metadata>
  <hpf:manifest>
    <hpf:item id="header" href="Contents/header.xml" media-type="application/xml"/>
    <hpf:item id="section0" href="Contents/section0.xml" media-type="application/xml"/>
    <hpf:item id="settings" href="settings.xml" media-type="application/xml"/>
  </hpf:manifest>
  <hpf:spine>
    <hpf:itemref idref="header"/>
    <hpf:itemref idref="section0"/>
  </hpf:spine>
</hpf:package>"""

# 7개 언어 폰트 매핑 (HANGUL / LATIN / HANJA / JAPANESE / OTHER / SYMBOL / USER)
_FONT_LANGS = ["HANGUL", "LATIN", "HANJA", "JAPANESE", "OTHER", "SYMBOL", "USER"]


def _header_xml() -> str:
    fontfaces = []
    for lang in _FONT_LANGS:
        fontfaces.append(
            f'<hh:fontface lang="{lang}" fontCnt="1">'
            f'<hh:font id="0" face="함초롬바탕" type="TTF" isEmbedded="0">'
            f'<hh:typeInfo familyType="FCAT_GOTHIC" weight="0" proportion="0" contrast="0" '
            f'strokeVariation="0" armStyle="0" letterform="0" midline="0" xHeight="0"/>'
            f'</hh:font></hh:fontface>'
        )
    # charPr: 7개 언어 모두 fontRef=0
    lang_refs = "".join(f'<hh:{tag} hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>'
                        for tag in ["fontRef", "ratio", "spacing", "relSz", "offset"])
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" version="1.4" secCnt="1">
<hh:beginNum page="1" footnote="1" endnote="1" pic="1" tbl="1" equation="1"/>
<hh:refList>
<hh:fontfaces itemCnt="7">{''.join(fontfaces)}</hh:fontfaces>
<hh:borderFills itemCnt="1">
<hh:borderFill id="1" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0">
<hh:slash type="NONE" Crooked="0" isCounter="0"/><hh:backSlash type="NONE" Crooked="0" isCounter="0"/>
<hh:leftBorder type="NONE" width="0.1 mm" color="#000000"/><hh:rightBorder type="NONE" width="0.1 mm" color="#000000"/>
<hh:topBorder type="NONE" width="0.1 mm" color="#000000"/><hh:bottomBorder type="NONE" width="0.1 mm" color="#000000"/>
<hh:diagonal type="SOLID" width="0.1 mm" color="#000000"/>
</hh:borderFill>
</hh:borderFills>
<hh:charProperties itemCnt="1">
<hh:charPr id="0" height="1000" textColor="#000000" shadeColor="none" useFontSpace="0" useKerning="0" symMark="NONE" borderFillIDRef="1">
{lang_refs}
<hh:bold/><hh:underline type="NONE" shape="SOLID" color="#000000"/>
<hh:strikeout shape="NONE" color="#000000"/><hh:outline type="NONE"/><hh:shadow type="NONE" color="#C0C0C0" offsetX="10" offsetY="10"/>
</hh:charPr>
</hh:charProperties>
<hh:tabProperties itemCnt="1">
<hh:tabPr id="0" autoTabLeft="0" autoTabRight="0"/>
</hh:tabProperties>
<hh:numberings itemCnt="1">
<hh:numbering id="1" start="0"><hh:paraHead start="1" level="1" align="LEFT" useInstWidth="1" autoIndent="1" widthAdjust="0" textOffsetType="PERCENT" textOffset="50" numFormat="DIGIT" charPrIDRef="4294967295">^1.</hh:paraHead></hh:numbering>
</hh:numberings>
<hh:paraProperties itemCnt="1">
<hh:paraPr id="0" tabPrIDRef="0" condense="0" fontLineHeight="0" snapToGrid="1" suppressLineNumbers="0" checked="0">
<hh:align horizontal="LEFT" vertical="BASELINE"/>
<hh:heading type="NONE" idRef="0" level="0"/>
<hh:breakSetting breakLatinWord="KEEP_WORD" breakNonLatinWord="KEEP_WORD" widowOrphan="0" keepWithNext="0" keepLines="0" pageBreakBefore="0" lineWrap="BREAK"/>
<hh:margin><hc:intent value="0" unit="HWPUNIT"/><hc:left value="0" unit="HWPUNIT"/><hc:right value="0" unit="HWPUNIT"/><hc:prev value="0" unit="HWPUNIT"/><hc:next value="0" unit="HWPUNIT"/></hh:margin>
<hh:lineSpacing type="PERCENT" value="160" unit="HWPUNIT"/>
<hh:border borderFillIDRef="1" offsetLeft="0" offsetRight="0" offsetTop="0" offsetBottom="0" connect="0" ignoreMargin="0"/>
</hh:paraPr>
</hh:paraProperties>
<hh:styles itemCnt="1">
<hh:style id="0" type="PARA" name="바탕글" engName="Normal" paraPrIDRef="0" charPrIDRef="0" nextStyleIDRef="0" langID="1042" lockForm="0"/>
</hh:styles>
</hh:refList>
</hh:head>"""


def _para(text: str, bold: bool = False) -> str:
    """단순 문단 한 개 생성. (charPr 0 사용, 굵게는 별도 미지원 — 단순화)"""
    return (
        f'<hp:p paraPrIDRef="0" styleIDRef="0">'
        f'<hp:run charPrIDRef="0"><hp:t>{_hx(text)}</hp:t></hp:run>'
        f'</hp:p>'
    )


def _section_xml(r: dict) -> str:
    paras = []
    # 첫 문단에 구역 정의(secPr) 포함
    first = (
        '<hp:p paraPrIDRef="0" styleIDRef="0"><hp:run charPrIDRef="0">'
        '<hp:secPr id="0" textDirection="HORIZONTAL" spaceColumns="0" tabStop="8000" tabStopVal="4000" tabStopUnit="HWPUNIT" outlineShapeIDRef="0" memoShapeIDRef="0" textVerticalWidthHead="0">'
        '<hp:grid lineGrid="0" charGrid="0" wonggojiFormat="0" strtnum="1"/>'
        '<hp:startNum pageStartsOn="BOTH" page="0" pic="0" tbl="0" equation="0"/>'
        '<hp:pagePr landscape="WIDELY" width="59528" height="84188" gutterType="LEFT_ONLY">'
        '<hp:margin header="4252" footer="4252" gutter="0" left="8504" right="8504" top="5668" bottom="4252"/>'
        '</hp:pagePr>'
        '</hp:secPr>'
        f'<hp:t>{_hx("일일 에너지 보고서")}</hp:t></hp:run></hp:p>'
    )
    paras.append(first)
    paras.append(_para(f"Honda R&D Europe GmbH — {r.get('date', '')}"))
    paras.append(_para(f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}"))
    paras.append(_para(""))
    paras.append(_para("[ KPI 요약 ]"))
    for label, value in _kpi_rows(r):
        paras.append(_para(f"  · {label}: {value}"))

    if r.get("ai_summary"):
        paras.append(_para(""))
        paras.append(_para("[ AI 요약 ]"))
        for line in r["ai_summary"].split("\n"):
            if line.strip():
                paras.append(_para("  " + line.strip()))

    hourly = _hourly_rows(r)
    if hourly:
        paras.append(_para(""))
        paras.append(_para("[ 시간대별 전력 프로파일 ]"))
        paras.append(_para("  " + " | ".join(_HOURLY_HEADER)))
        for row in hourly:
            paras.append(_para("  " + " | ".join(row)))

    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" '
        'xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph" '
        'xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">'
        + "".join(paras) +
        '</hs:sec>'
    )


def build_hwpx(r: dict) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        # mimetype 은 반드시 첫 엔트리 + 무압축
        zi = zipfile.ZipInfo("mimetype")
        zi.compress_type = zipfile.ZIP_STORED
        z.writestr(zi, "application/hwp+zip")

        z.writestr("version.xml", _VERSION_XML)
        z.writestr("settings.xml", _SETTINGS_XML)
        z.writestr("Contents/content.hpf", _CONTENT_HPF)
        z.writestr("Contents/header.xml", _header_xml())
        z.writestr("Contents/section0.xml", _section_xml(r))
        z.writestr("META-INF/container.xml", _CONTAINER_XML)
        z.writestr("META-INF/manifest.xml", _MANIFEST_XML)
    return buf.getvalue()


# ── 디스패처 ─────────────────────────────────────────────────────

_BUILDERS = {"pdf": build_pdf, "docx": build_docx, "hwpx": build_hwpx}
_MEDIA = {
    "pdf":  "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "hwpx": "application/hwp+zip",
}


def render(report: dict, fmt: str) -> tuple[bytes, str, str]:
    """(bytes, media_type, filename) 반환. 지원하지 않는 포맷이면 ValueError."""
    fmt = fmt.lower()
    if fmt not in _BUILDERS:
        raise ValueError(f"지원하지 않는 포맷: {fmt}")
    data = _BUILDERS[fmt](report)
    filename = f"daily_report_{report.get('date', 'report')}.{fmt}"
    return data, _MEDIA[fmt], filename
